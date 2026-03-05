import os
import sys
import time
import subprocess
import re
import shutil
import urllib.request
import ctypes
import gc
import json
import threading

try:
    ctypes.windll.shcore.SetProcessDpiAwareness(1)
except Exception:
    ctypes.windll.user32.SetProcessDPIAware()

import warnings
warnings.filterwarnings("ignore")

# ── PyQt6 availability check ──────────────────────────────────────────────────
try:
    from PyQt6.QtWidgets import (
        QApplication, QMainWindow, QWidget, QFrame,
        QVBoxLayout, QHBoxLayout, QGridLayout,
        QLabel, QLineEdit, QPushButton, QComboBox, QCheckBox,
        QSpinBox, QProgressBar, QPlainTextEdit, QSizeGrip,
        QFileDialog, QMessageBox, QGraphicsOpacityEffect
    )
    from PyQt6.QtCore import (
        Qt, QThread, pyqtSignal, QTimer, QUrl,
        QEasingCurve, QPropertyAnimation, QPoint
    )
    from PyQt6.QtGui import QDesktopServices, QCursor
    PYQT6_OK = True
except ImportError:
    PYQT6_OK = False

# ── Constants ─────────────────────────────────────────────────────────────────
APP_VERSION     = "1.3.1"
GITHUB_REPO     = "acrilot/whisper-gui"
GITHUB_PAGE_URL = "https://github.com/acrilot/whisper-gui"
FFMPEG_BAT_URL  = (
    "https://raw.githubusercontent.com/acrilot/whisper-gui"
    "/refs/heads/main/install_ffmpeg.bat"
)

BASE_DIR = os.getcwd()

MODEL_PATHS = {
    "faster_large": os.path.join(BASE_DIR, "faster_whisper_large_v3"),
    "faster_turbo": os.path.join(BASE_DIR, "faster_whisper_turbo"),
    "std_large":    os.path.join(BASE_DIR, "large-v3.pt"),
    "std_turbo":    os.path.join(BASE_DIR, "large-v3-turbo.pt"),
    "std_medium":   os.path.join(BASE_DIR, "medium.pt"),
}

MODEL_OPTIONS = {
    "Faster Whisper — Large V3 (Best Quality)": "faster_large",
    "Faster Whisper — Turbo (Fastest)":         "faster_turbo",
    "Standard Whisper — Large V3":              "std_large",
    "Standard Whisper — Turbo":                 "std_turbo",
    "Standard Whisper — Medium":                "std_medium",
}

# Language codes by index — order must match STRINGS["en"]["lang_names"]
LANGUAGE_CODES   = [None, "tr", "en", "de", "fr", "es", "zh", "ja"]
FORMAT_OPTIONS   = ["TXT", "PDF", "DOCX", "XML (Subtitle)"]
COMPUTE_OPTIONS  = ["int8", "float16", "int8_float16", "float32"]

LIBS_TO_INSTALL = [
    "PyQt6", "openai-whisper", "faster-whisper",
    "huggingface_hub", "fpdf", "python-docx",
]
PYTORCH_INSTALL_ARGS = [
    "torch", "torchvision", "torchaudio",
    "--index-url", "https://download.pytorch.org/whl/cu124",
]

# ── UI strings (TR / EN / DE) ─────────────────────────────────────────────────
STRINGS = {
    "en": {
        "window_title":      "Whisper Transcript",
        "header_sub":        "  Transcript Extractor",
        "sec_files":         "FILES & FORMAT",
        "sec_model":         "MODEL & SETTINGS",
        "sec_options":       "OPTIONS",
        "sec_progress":      "PROGRESS",
        "sec_log":           "TRANSCRIPT LOG",
        "lbl_language":      "Language",
        "lbl_compute":       "Compute",
        "lbl_beam":          "Beam",
        "lbl_output":        "Output",
        "ph_files":          "Select audio or video files...",
        "ph_output":         "Output path...",
        "btn_browse":        "Browse",
        "btn_start":         "Start",
        "btn_stop":          "Stop and Save",
        "btn_clear":         "Clear",
        "chk_vad":           "Anti-Loop (VAD)",
        "chk_single":        "Single Block",
        "chk_timestamps":    "Timestamps",
        "chk_translate":     "Translate to English",
        "chk_install":       "Install Mode",
        "status_ready":      "Ready",
        "status_done":       "Done",
        "status_stopped":    "Stopped",
        "status_error":      "Error",
        "lang_names": [
            "Auto Detect", "Turkish", "English", "German",
            "French", "Spanish", "Chinese", "Japanese",
        ],
        "ffmpeg_title":      "FFmpeg Missing",
        "ffmpeg_msg":        (
            "FFmpeg was not found on your system. "
            "It is required for audio processing.\n\n"
            "Download and install it automatically from GitHub?"
        ),
        "ffmpeg_warn":       "FFmpeg is missing. The application may not work correctly.",
        "ffmpeg_restart":    "FFmpeg installed. The application will restart.",
        "stop_title":        "Stop",
        "stop_msg":          "Stop the process and save progress so far?",
        "exit_title":        "Exit",
        "exit_msg":          "A transcription is in progress. Exit anyway?",
        "done_title":        "Complete",
        "done_msg":          "All files processed.\nTotal time: {:.2f}s",
        "update_title":      "Update Available",
        "update_msg":        (
            "A new version is available: V{new}\n"
            "Current version: V{cur}\n\n"
            "Download and restart now?"
        ),
        "update_no_asset":   "No .pyw file found for this release. Check GitHub manually.",
        "update_done":       "V{new} downloaded.\nPrevious version backed up as .bak\nRestarting...",
        "update_fail":       "Update failed:\n{err}\n\nPlease download manually from GitHub.",
        "batch_output_info": "In batch mode, outputs are saved next to each source file.",
        "warn_no_file":      "Please select at least one file.",
        "install_error":     "{lib} installation failed.\nDetail: {detail}",
        "pytorch_error":     "PyTorch (CUDA) installation failed.",
    },
    "tr": {
        "window_title":      "Whisper Transkript",
        "header_sub":        "  Transkript Çıkarıcı",
        "sec_files":         "DOSYA & FORMAT",
        "sec_model":         "MODEL & AYARLAR",
        "sec_options":       "SEÇENEKLER",
        "sec_progress":      "İLERLEME",
        "sec_log":           "TRANSKRİPT KAYITLARI",
        "lbl_language":      "Dil",
        "lbl_compute":       "Compute",
        "lbl_beam":          "Beam",
        "lbl_output":        "Çıktı",
        "ph_files":          "Ses veya video dosyası seçin...",
        "ph_output":         "Çıktı yolu...",
        "btn_browse":        "Seç",
        "btn_start":         "Başlat",
        "btn_stop":          "Durdur ve Kaydet",
        "btn_clear":         "Temizle",
        "chk_vad":           "Anti-Loop (VAD)",
        "chk_single":        "Tek Blok Metin",
        "chk_timestamps":    "Zaman Damgası",
        "chk_translate":     "İngilizce'ye Çevir",
        "chk_install":       "Kurulum Modu",
        "status_ready":      "Hazır",
        "status_done":       "Tamamlandı",
        "status_stopped":    "Durduruldu",
        "status_error":      "Hata",
        "lang_names": [
            "Otomatik", "Türkçe", "İngilizce", "Almanca",
            "Fransızca", "İspanyolca", "Çince", "Japonca",
        ],
        "ffmpeg_title":      "FFmpeg Eksik",
        "ffmpeg_msg":        (
            "Sisteminizde FFmpeg bulunamadı. "
            "Ses işleme için gereklidir.\n\n"
            "GitHub üzerinden otomatik kurulum yapılsın mı?"
        ),
        "ffmpeg_warn":       "FFmpeg eksik. Uygulama düzgün çalışmayabilir.",
        "ffmpeg_restart":    "FFmpeg kuruldu. Uygulama yeniden başlatılıyor.",
        "stop_title":        "Durdur",
        "stop_msg":          "İşlem durdurulsun ve kaydedilsin mi?",
        "exit_title":        "Çıkış",
        "exit_msg":          "Devam eden bir transkript işlemi var. Yine de çıkılsın mı?",
        "done_title":        "Tamamlandı",
        "done_msg":          "Tüm dosyalar işlendi.\nToplam süre: {:.2f}s",
        "update_title":      "Güncelleme Mevcut",
        "update_msg":        (
            "Yeni sürüm mevcut: V{new}\n"
            "Mevcut sürüm: V{cur}\n\n"
            "İndirilip yeniden başlatılsın mı?"
        ),
        "update_no_asset":   "Bu sürüm için .pyw dosyası bulunamadı. GitHub'ı manuel kontrol edin.",
        "update_done":       "V{new} indirildi.\nEski sürüm .bak olarak yedeklendi.\nYeniden başlatılıyor...",
        "update_fail":       "Güncelleme başarısız:\n{err}\n\nLütfen GitHub'dan manuel indirin.",
        "batch_output_info": "Toplu işlemde çıktılar kaynak dosyaların yanına kaydedilir.",
        "warn_no_file":      "Lütfen en az bir dosya seçin.",
        "install_error":     "{lib} kurulumu başarısız.\nDetay: {detail}",
        "pytorch_error":     "PyTorch (CUDA) kurulumu başarısız.",
    },
    "de": {
        "window_title":      "Whisper Transkript",
        "header_sub":        "  Transkript-Extraktor",
        "sec_files":         "DATEIEN & FORMAT",
        "sec_model":         "MODELL & EINSTELLUNGEN",
        "sec_options":       "OPTIONEN",
        "sec_progress":      "FORTSCHRITT",
        "sec_log":           "TRANSKRIPT-PROTOKOLL",
        "lbl_language":      "Sprache",
        "lbl_compute":       "Rechenleistung",
        "lbl_beam":          "Beam",
        "lbl_output":        "Ausgabe",
        "ph_files":          "Audio- oder Videodateien auswählen...",
        "ph_output":         "Ausgabepfad...",
        "btn_browse":        "Durchsuchen",
        "btn_start":         "Starten",
        "btn_stop":          "Stoppen & Speichern",
        "btn_clear":         "Leeren",
        "chk_vad":           "Anti-Schleife (VAD)",
        "chk_single":        "Einzelblock",
        "chk_timestamps":    "Zeitstempel",
        "chk_translate":     "Ins Englische übersetzen",
        "chk_install":       "Installationsmodus",
        "status_ready":      "Bereit",
        "status_done":       "Fertig",
        "status_stopped":    "Gestoppt",
        "status_error":      "Fehler",
        "lang_names": [
            "Automatisch erkennen", "Türkisch", "Englisch", "Deutsch",
            "Französisch", "Spanisch", "Chinesisch", "Japanisch"
        ],
        "ffmpeg_title":      "FFmpeg fehlt",
        "ffmpeg_msg":        (
            "FFmpeg wurde auf Ihrem System nicht gefunden. "
            "Es wird für die Audioverarbeitung benötigt.\n\n"
            "Automatisch von GitHub herunterladen und installieren?"
        ),
        "ffmpeg_warn":       "FFmpeg fehlt. Die Anwendung funktioniert möglicherweise nicht korrekt.",
        "ffmpeg_restart":    "FFmpeg installiert. Die Anwendung wird neu gestartet.",
        "stop_title":        "Stoppen",
        "stop_msg":          "Vorgang abbrechen und bisherigen Fortschritt speichern?",
        "exit_title":        "Beenden",
        "exit_msg":          "Eine Transkription ist im Gange. Trotzdem beenden?",
        "done_title":        "Abgeschlossen",
        "done_msg":          "Alle Dateien verarbeitet.\nGesamtzeit: {:.2f}s",
        "update_title":      "Update verfügbar",
        "update_msg":        (
            "Eine neue Version ist verfügbar: V{new}\n"
            "Aktuelle Version: V{cur}\n\n"
            "Jetzt herunterladen und neu starten?"
        ),
        "update_no_asset":   "Keine .pyw-Datei für dieses Release gefunden. Bitte manuell auf GitHub prüfen.",
        "update_done":       "V{new} heruntergeladen.\nVorherige Version als .bak gesichert.\nNeustart...",
        "update_fail":       "Update fehlgeschlagen:\n{err}\n\nBitte manuell von GitHub herunterladen.",
        "batch_output_info": "Im Batch-Modus werden die Ausgaben neben der jeweiligen Quelldatei gespeichert.",
        "warn_no_file":      "Bitte wählen Sie mindestens eine Datei aus.",
        "install_error":     "Installation von {lib} fehlgeschlagen.\nDetails: {detail}",
        "pytorch_error":     "PyTorch (CUDA) Installation fehlgeschlagen."
    },
    "es": {
        "window_title":      "Transcriptor Whisper",
        "header_sub":        "  Extractor de Transcripciones",
        "sec_files":         "ARCHIVOS Y FORMATO",
        "sec_model":         "MODELO Y AJUSTES",
        "sec_options":       "OPCIONES",
        "sec_progress":      "PROGRESO",
        "sec_log":           "REGISTRO DE TRANSCRIPCIÓN",
        "lbl_language":      "Idioma",
        "lbl_compute":       "Cómputo",
        "lbl_beam":          "Beam",
        "lbl_output":        "Salida",
        "ph_files":          "Seleccionar archivos de audio o video...",
        "ph_output":         "Ruta de salida...",
        "btn_browse":        "Explorar",
        "btn_start":         "Iniciar",
        "btn_stop":          "Detener y Guardar",
        "btn_clear":         "Limpiar",
        "chk_vad":           "Anti-Bucle (VAD)",
        "chk_single":        "Bloque Único",
        "chk_timestamps":    "Marcas de tiempo",
        "chk_translate":     "Traducir al Inglés",
        "chk_install":       "Modo de Instalación",
        "status_ready":      "Listo",
        "status_done":       "Hecho",
        "status_stopped":    "Detenido",
        "status_error":      "Error",
        "lang_names": [
            "Autodetectar", "Turco", "Inglés", "Alemán",
            "Francés", "Español", "Chino", "Japonés"
        ],
        "ffmpeg_title":      "Falta FFmpeg",
        "ffmpeg_msg":        "No se encontró FFmpeg en su sistema. Es necesario para el procesamiento de audio.\n\n¿Descargarlo e instalarlo automáticamente desde GitHub?",
        "ffmpeg_warn":       "Falta FFmpeg. La aplicación podría no funcionar correctamente.",
        "ffmpeg_restart":    "FFmpeg instalado. La aplicación se reiniciará.",
        "stop_title":        "Detener",
        "stop_msg":          "¿Detener el proceso y guardar el progreso hasta ahora?",
        "exit_title":        "Salir",
        "exit_msg":          "Hay una transcripción en curso. ¿Salir de todos modos?",
        "done_title":        "Completado",
        "done_msg":          "Todos los archivos procesados.\nTiempo total: {:.2f}s",
        "update_title":      "Actualización Disponible",
        "update_msg":        "Hay una nueva versión disponible: V{new}\nVersión actual: V{cur}\n\n¿Descargar y reiniciar ahora?",
        "update_no_asset":   "No se encontró el archivo .pyw para esta versión. Verifique GitHub manualmente.",
        "update_done":       "V{new} descargada.\nVersión anterior respaldada como .bak\nReiniciando...",
        "update_fail":       "Falló la actualización:\n{err}\n\nPor favor, descargue manualmente desde GitHub.",
        "batch_output_info": "En modo por lotes, las salidas se guardan junto a cada archivo de origen.",
        "warn_no_file":      "Por favor, seleccione al menos un archivo.",
        "install_error":     "Falló la instalación de {lib}.\nDetalle: {detail}",
        "pytorch_error":     "Falló la instalación de PyTorch (CUDA)."
    },
    "pt": {
        "window_title":      "Transcritor Whisper",
        "header_sub":        "  Extrator de Transcrições",
        "sec_files":         "ARQUIVOS E FORMATO",
        "sec_model":         "MODELO E CONFIGURAÇÕES",
        "sec_options":       "OPÇÕES",
        "sec_progress":      "PROGRESSO",
        "sec_log":           "REGISTRO DE TRANSCRIÇÃO",
        "lbl_language":      "Idioma",
        "lbl_compute":       "Computação",
        "lbl_beam":          "Beam",
        "lbl_output":        "Saída",
        "ph_files":          "Selecione arquivos de áudio ou vídeo...",
        "ph_output":         "Caminho de saída...",
        "btn_browse":        "Procurar",
        "btn_start":         "Iniciar",
        "btn_stop":          "Parar e Salvar",
        "btn_clear":         "Limpar",
        "chk_vad":           "Anti-Loop (VAD)",
        "chk_single":        "Bloco Único",
        "chk_timestamps":    "Carimbos de tempo",
        "chk_translate":     "Traduzir para Inglês",
        "chk_install":       "Modo de Instalação",
        "status_ready":      "Pronto",
        "status_done":       "Concluído",
        "status_stopped":    "Parado",
        "status_error":      "Erro",
        "lang_names": [
            "Detecção Automática", "Turco", "Inglês", "Alemão",
            "Francês", "Espanhol", "Chinês", "Japonês"
        ],
        "ffmpeg_title":      "FFmpeg Ausente",
        "ffmpeg_msg":        "O FFmpeg não foi encontrado no seu sistema. Ele é necessário para o processamento de áudio.\n\nBaixar e instalar automaticamente do GitHub?",
        "ffmpeg_warn":       "FFmpeg está ausente. O aplicativo pode não funcionar corretamente.",
        "ffmpeg_restart":    "FFmpeg instalado. O aplicativo será reiniciado.",
        "stop_title":        "Parar",
        "stop_msg":          "Parar o processo e salvar o progresso até agora?",
        "exit_title":        "Sair",
        "exit_msg":          "Uma transcrição está em andamento. Sair de qualquer maneira?",
        "done_title":        "Completo",
        "done_msg":          "Todos os arquivos processados.\nTempo total: {:.2f}s",
        "update_title":      "Atualização Disponível",
        "update_msg":        "Uma nova versão está disponível: V{new}\nVersão atual: V{cur}\n\nBaixar e reiniciar agora?",
        "update_no_asset":   "Nenhum arquivo .pyw encontrado para este lançamento. Verifique o GitHub manualmente.",
        "update_done":       "V{new} baixada.\nVersão anterior salva como .bak\nReiniciando...",
        "update_fail":       "A atualização falhou:\n{err}\n\nPor favor, baixe manualmente do GitHub.",
        "batch_output_info": "No modo em lote, as saídas são salvas ao lado de cada arquivo de origem.",
        "warn_no_file":      "Por favor, selecione pelo menos um arquivo.",
        "install_error":     "A instalação de {lib} falhou.\nDetalhe: {detail}",
        "pytorch_error":     "A instalação do PyTorch (CUDA) falhou."
    },
    "ja": {
        "window_title":      "Whisper 文字起こし",
        "header_sub":        "  文字起こし抽出ツール",
        "sec_files":         "ファイルとフォーマット",
        "sec_model":         "モデルと設定",
        "sec_options":       "オプション",
        "sec_progress":      "進捗",
        "sec_log":           "文字起こしログ",
        "lbl_language":      "言語",
        "lbl_compute":       "計算処理",
        "lbl_beam":          "ビーム",
        "lbl_output":        "出力",
        "ph_files":          "音声または動画ファイルを選択...",
        "ph_output":         "出力先パス...",
        "btn_browse":        "参照",
        "btn_start":         "開始",
        "btn_stop":          "停止して保存",
        "btn_clear":         "クリア",
        "chk_vad":           "ループ防止 (VAD)",
        "chk_single":        "シングルブロック",
        "chk_timestamps":    "タイムスタンプ",
        "chk_translate":     "英語に翻訳",
        "chk_install":       "インストールモード",
        "status_ready":      "準備完了",
        "status_done":       "完了",
        "status_stopped":    "停止",
        "status_error":      "エラー",
        "lang_names": [
            "自動検出", "トルコ語", "英語", "ドイツ語",
            "フランス語", "スペイン語", "中国語", "日本語"
        ],
        "ffmpeg_title":      "FFmpeg が見つかりません",
        "ffmpeg_msg":        "システムに FFmpeg が見つかりませんでした。音声処理に必要です。\n\nGitHub から自動的にダウンロードしてインストールしますか？",
        "ffmpeg_warn":       "FFmpeg がありません。アプリケーションが正常に動作しない可能性があります。",
        "ffmpeg_restart":    "FFmpeg がインストールされました。アプリケーションを再起動します。",
        "stop_title":        "停止",
        "stop_msg":          "プロセスを停止し、これまでの進捗を保存しますか？",
        "exit_title":        "終了",
        "exit_msg":          "文字起こしが進行中です。それでも終了しますか？",
        "done_title":        "完了",
        "done_msg":          "すべてのファイルが処理されました。\n合計時間: {:.2f}秒",
        "update_title":      "アップデート利用可能",
        "update_msg":        "新しいバージョンが利用可能です: V{new}\n現在のバージョン: V{cur}\n\n今すぐダウンロードして再起動しますか？",
        "update_no_asset":   "このリリースの .pyw ファイルが見つかりません。GitHub を手動で確認してください。",
        "update_done":       "V{new} がダウンロードされました。\n前のバージョンは .bak としてバックアップされました\n再起動しています...",
        "update_fail":       "アップデートに失敗しました:\n{err}\n\nGitHub から手動でダウンロードしてください。",
        "batch_output_info": "バッチモードでは、出力は各ソースファイルと同じ場所に保存されます。",
        "warn_no_file":      "少なくとも1つのファイルを選択してください。",
        "install_error":     "{lib} のインストールに失敗しました。\n詳細: {detail}",
        "pytorch_error":     "PyTorch (CUDA) のインストールに失敗しました。"
    },
    "zh": {
        "window_title":      "Whisper 转录器",
        "header_sub":        "  转录提取工具",
        "sec_files":         "文件与格式",
        "sec_model":         "模型与设置",
        "sec_options":       "选项",
        "sec_progress":      "进度",
        "sec_log":           "转录日志",
        "lbl_language":      "语言",
        "lbl_compute":       "计算",
        "lbl_beam":          "Beam",
        "lbl_output":        "输出",
        "ph_files":          "选择音频或视频文件...",
        "ph_output":         "输出路径...",
        "btn_browse":        "浏览",
        "btn_start":         "开始",
        "btn_stop":          "停止并保存",
        "btn_clear":         "清除",
        "chk_vad":           "防循环 (VAD)",
        "chk_single":        "单区块",
        "chk_timestamps":    "时间戳",
        "chk_translate":     "翻译为英语",
        "chk_install":       "安装模式",
        "status_ready":      "就绪",
        "status_done":       "完成",
        "status_stopped":    "已停止",
        "status_error":      "错误",
        "lang_names": [
            "自动检测", "土耳其语", "英语", "德语",
            "法语", "西班牙语", "中文", "日语"
        ],
        "ffmpeg_title":      "缺少 FFmpeg",
        "ffmpeg_msg":        "您的系统中未找到 FFmpeg。音频处理需要它。\n\n是否从 GitHub 自动下载并安装？",
        "ffmpeg_warn":       "缺少 FFmpeg。应用程序可能无法正常工作。",
        "ffmpeg_restart":    "FFmpeg 已安装。应用程序将重启。",
        "stop_title":        "停止",
        "stop_msg":          "是否停止进程并保存当前的进度？",
        "exit_title":        "退出",
        "exit_msg":          "转录正在进行中。确认要退出吗？",
        "done_title":        "完成",
        "done_msg":          "所有文件处理完毕。\n总用时: {:.2f}秒",
        "update_title":      "可用更新",
        "update_msg":        "有新版本可用: V{new}\n当前版本: V{cur}\n\n是否立即下载并重启？",
        "update_no_asset":   "未找到此版本的 .pyw 文件。请手动检查 GitHub。",
        "update_done":       "V{new} 已下载。\n旧版本已备份为 .bak\n正在重启...",
        "update_fail":       "更新失败:\n{err}\n\n请从 GitHub 手动下载。",
        "batch_output_info": "在批量模式下，输出文件将保存在每个源文件的旁边。",
        "warn_no_file":      "请至少选择一个文件。",
        "install_error":     "{lib} 安装失败。\n详细信息: {detail}",
        "pytorch_error":     "PyTorch (CUDA) 安装失败。"
    },
    "fr": {
        "window_title":      "Transcripteur Whisper",
        "header_sub":        "  Extracteur de Transcription",
        "sec_files":         "FICHIERS ET FORMAT",
        "sec_model":         "MODÈLE ET PARAMÈTRES",
        "sec_options":       "OPTIONS",
        "sec_progress":      "PROGRESSION",
        "sec_log":           "JOURNAL DE TRANSCRIPTION",
        "lbl_language":      "Langue",
        "lbl_compute":       "Calcul",
        "lbl_beam":          "Faisceau",
        "lbl_output":        "Sortie",
        "ph_files":          "Sélectionnez des fichiers audio ou vidéo...",
        "ph_output":         "Chemin de sortie...",
        "btn_browse":        "Parcourir",
        "btn_start":         "Démarrer",
        "btn_stop":          "Arrêter et Enregistrer",
        "btn_clear":         "Effacer",
        "chk_vad":           "Anti-Boucle (VAD)",
        "chk_single":        "Bloc Unique",
        "chk_timestamps":    "Horodatages",
        "chk_translate":     "Traduire en anglais",
        "chk_install":       "Mode d'Installation",
        "status_ready":      "Prêt",
        "status_done":       "Terminé",
        "status_stopped":    "Arrêté",
        "status_error":      "Erreur",
        "lang_names": [
            "Détection Auto", "Turc", "Anglais", "Allemand",
            "Français", "Espagnol", "Chinois", "Japonais"
        ],
        "ffmpeg_title":      "FFmpeg Manquant",
        "ffmpeg_msg":        "FFmpeg est introuvable sur votre système. Il est requis pour le traitement audio.\n\nTélécharger et installer automatiquement depuis GitHub ?",
        "ffmpeg_warn":       "FFmpeg est manquant. L'application risque de ne pas fonctionner correctement.",
        "ffmpeg_restart":    "FFmpeg installé. L'application va redémarrer.",
        "stop_title":        "Arrêter",
        "stop_msg":          "Arrêter le processus et enregistrer la progression actuelle ?",
        "exit_title":        "Quitter",
        "exit_msg":          "Une transcription est en cours. Quitter quand même ?",
        "done_title":        "Terminé",
        "done_msg":          "Tous les fichiers ont été traités.\nTemps total : {:.2f}s",
        "update_title":      "Mise à jour Disponible",
        "update_msg":        "Une nouvelle version est disponible : V{new}\nVersion actuelle : V{cur}\n\nTélécharger et redémarrer maintenant ?",
        "update_no_asset":   "Aucun fichier .pyw trouvé pour cette version. Vérifiez manuellement sur GitHub.",
        "update_done":       "V{new} téléchargée.\nVersion précédente sauvegardée en .bak\nRedémarrage...",
        "update_fail":       "Échec de la mise à jour :\n{err}\n\nVeuillez télécharger manuellement depuis GitHub.",
        "batch_output_info": "En mode traitement par lots, les sorties sont enregistrées à côté de chaque fichier source.",
        "warn_no_file":      "Veuillez sélectionner au moins un fichier.",
        "install_error":     "L'installation de {lib} a échoué.\nDétail : {detail}",
        "pytorch_error":     "L'installation de PyTorch (CUDA) a échoué."
    },
    "it": {
        "window_title":      "Trascrittore Whisper",
        "header_sub":        "  Estrattore di Trascrizioni",
        "sec_files":         "FILE E FORMATO",
        "sec_model":         "MODELLO E IMPOSTAZIONI",
        "sec_options":       "OPZIONI",
        "sec_progress":      "PROGRESSO",
        "sec_log":           "LOG DI TRASCRIZIONE",
        "lbl_language":      "Lingua",
        "lbl_compute":       "Calcolo",
        "lbl_beam":          "Beam",
        "lbl_output":        "Output",
        "ph_files":          "Seleziona file audio o video...",
        "ph_output":         "Percorso di output...",
        "btn_browse":        "Sfoglia",
        "btn_start":         "Avvia",
        "btn_stop":          "Ferma e Salva",
        "btn_clear":         "Pulisci",
        "chk_vad":           "Anti-Loop (VAD)",
        "chk_single":        "Blocco Singolo",
        "chk_timestamps":    "Marche temporali",
        "chk_translate":     "Traduci in Inglese",
        "chk_install":       "Modalità Installazione",
        "status_ready":      "Pronto",
        "status_done":       "Fatto",
        "status_stopped":    "Fermato",
        "status_error":      "Errore",
        "lang_names": [
            "Rilevamento Auto", "Turco", "Inglese", "Tedesco",
            "Francese", "Spagnolo", "Cinese", "Giapponese"
        ],
        "ffmpeg_title":      "FFmpeg Mancante",
        "ffmpeg_msg":        "FFmpeg non è stato trovato sul tuo sistema. È richiesto per l'elaborazione audio.\n\nScaricare e installare automaticamente da GitHub?",
        "ffmpeg_warn":       "FFmpeg è mancante. L'applicazione potrebbe non funzionare correttamente.",
        "ffmpeg_restart":    "FFmpeg installato. L'applicazione verrà riavviata.",
        "stop_title":        "Ferma",
        "stop_msg":          "Fermare il processo e salvare i progressi finora?",
        "exit_title":        "Esci",
        "exit_msg":          "Una trascrizione è in corso. Uscire comunque?",
        "done_title":        "Completato",
        "done_msg":          "Tutti i file processati.\nTempo totale: {:.2f}s",
        "update_title":      "Aggiornamento Disponibile",
        "update_msg":        "È disponibile una nuova versione: V{new}\nVersione attuale: V{cur}\n\nScaricare e riavviare ora?",
        "update_no_asset":   "Nessun file .pyw trovato per questa versione. Controlla GitHub manualmente.",
        "update_done":       "V{new} scaricata.\nVersione precedente salvata come .bak\nRiavvio in corso...",
        "update_fail":       "Aggiornamento fallito:\n{err}\n\nSi prega di scaricare manualmente da GitHub.",
        "batch_output_info": "In modalità batch, gli output vengono salvati accanto a ciascun file sorgente.",
        "warn_no_file":      "Si prega di selezionare almeno un file.",
        "install_error":     "Installazione di {lib} fallita.\nDettaglio: {detail}",
        "pytorch_error":     "Installazione di PyTorch (CUDA) fallita."
    },
    "ru": {
        "window_title":      "Транскрибатор Whisper",
        "header_sub":        "  Экстрактор Транскрипций",
        "sec_files":         "ФАЙЛЫ И ФОРМАТ",
        "sec_model":         "МОДЕЛЬ И НАСТРОЙКИ",
        "sec_options":       "ОПЦИИ",
        "sec_progress":      "ПРОГРЕСС",
        "sec_log":           "ЖУРНАЛ ТРАНСКРИПЦИИ",
        "lbl_language":      "Язык",
        "lbl_compute":       "Вычисления",
        "lbl_beam":          "Beam",
        "lbl_output":        "Вывод",
        "ph_files":          "Выберите аудио или видео файлы...",
        "ph_output":         "Путь для сохранения...",
        "btn_browse":        "Обзор",
        "btn_start":         "Старт",
        "btn_stop":          "Остановить и Сохранить",
        "btn_clear":         "Очистить",
        "chk_vad":           "Анти-цикл (VAD)",
        "chk_single":        "Один блок",
        "chk_timestamps":    "Таймкоды",
        "chk_translate":     "Перевести на английский",
        "chk_install":       "Режим установки",
        "status_ready":      "Готово",
        "status_done":       "Завершено",
        "status_stopped":    "Остановлено",
        "status_error":      "Ошибка",
        "lang_names": [
            "Автоопределение", "Турецкий", "Английский", "Немецкий",
            "Французский", "Испанский", "Китайский", "Японский"
        ],
        "ffmpeg_title":      "FFmpeg не найден",
        "ffmpeg_msg":        "FFmpeg не найден в вашей системе. Он необходим для обработки аудио.\n\nСкачать и установить его автоматически с GitHub?",
        "ffmpeg_warn":       "FFmpeg отсутствует. Приложение может работать некорректно.",
        "ffmpeg_restart":    "FFmpeg установлен. Приложение будет перезапущено.",
        "stop_title":        "Стоп",
        "stop_msg":          "Остановить процесс и сохранить текущий прогресс?",
        "exit_title":        "Выход",
        "exit_msg":          "Транскрибация еще не завершена. Выйти все равно?",
        "done_title":        "Завершено",
        "done_msg":          "Все файлы обработаны.\nОбщее время: {:.2f}с",
        "update_title":      "Доступно обновление",
        "update_msg":        "Доступна новая версия: V{new}\nТекущая версия: V{cur}\n\nСкачать и перезапустить сейчас?",
        "update_no_asset":   "Для этого релиза не найден файл .pyw. Проверьте GitHub вручную.",
        "update_done":       "V{new} скачана.\nПредыдущая версия сохранена как .bak\nПерезапуск...",
        "update_fail":       "Ошибка обновления:\n{err}\n\nПожалуйста, скачайте вручную с GitHub.",
        "batch_output_info": "В пакетном режиме файлы сохраняются рядом с каждым исходным файлом.",
        "warn_no_file":      "Пожалуйста, выберите хотя бы один файл.",
        "install_error":     "Ошибка установки {lib}.\nДетали: {detail}",
        "pytorch_error":     "Ошибка установки PyTorch (CUDA)."
    }
}

# ── Theme palette definitions ─────────────────────────────────────────────────
THEME_ACCENTS = {
    "yellow": {"accent": "#f0a500", "accent_hover": "#f5b832", "accent_press": "#d49200"},
    "red":    {"accent": "#ef4444", "accent_hover": "#f87171", "accent_press": "#dc2626"},
    "green":  {"accent": "#22c55e", "accent_hover": "#4ade80", "accent_press": "#16a34a"},
    "blue":   {"accent": "#3b82f6", "accent_hover": "#60a5fa", "accent_press": "#2563eb"},
}

DARK_PALETTE = {
    "bg_main":    "#0d1117",
    "bg_card":    "#161b22",
    "bg_input":   "#0d1117",
    "bg_subtle":  "#21262d",
    "border":     "#30363d",
    "border_sub": "#21262d",
    "text_main":  "#c9d1d9",
    "text_muted": "#8b949e",
    "text_dim":   "#484f58",
    "log_bg":     "#0d1117",
    "log_text":   "#8b949e",
}

LIGHT_PALETTE = {
    "bg_main":    "#f6f8fa",
    "bg_card":    "#ffffff",
    "bg_input":   "#f6f8fa",
    "bg_subtle":  "#eaeef2",
    "border":     "#d0d7de",
    "border_sub": "#eaeef2",
    "text_main":  "#1f2328",
    "text_muted": "#57606a",
    "text_dim":   "#9198a1",
    "log_bg":     "#f6f8fa",
    "log_text":   "#57606a",
}


def build_stylesheet(theme_name: str, dark: bool) -> str:
    acc   = THEME_ACCENTS[theme_name]
    base  = DARK_PALETTE if dark else LIGHT_PALETTE

    a        = acc["accent"]
    a_hover  = acc["accent_hover"]
    a_press  = acc["accent_press"]

    # Disabled button tones — always derived from accent at low opacity feel
    a_dis_bg = base["bg_subtle"]
    a_dis_fg = base["text_dim"]

    bg       = base["bg_main"]
    card     = base["bg_card"]
    inp      = base["bg_input"]
    subtle   = base["bg_subtle"]
    brd      = base["border"]
    brd_sub  = base["border_sub"]
    txt      = base["text_main"]
    muted    = base["text_muted"]
    dim      = base["text_dim"]
    log_bg   = base["log_bg"]
    log_txt  = base["log_text"]

    # Danger button colors (always reddish regardless of theme)
    if dark:
        danger_fg     = "#f87171"
        danger_brd    = "#3f1212"
        danger_hover  = a_hover
        danger_dis_fg = a
        danger_dis_brd= a
    else:
        danger_fg     = "#dc2626"
        danger_brd    = "#fca5a5"
        danger_hover  = a_hover
        danger_dis_fg = a
        danger_dis_brd= a

    # Input text color when disabled
    inp_dis_txt = dim
    inp_dis_brd = subtle
    muted_svg = muted.replace("#", "%23")
    font_family = "'Segoe UI Variable Display', 'Inter', 'Roboto', 'Segoe UI', sans-serif"

    return f"""
/* ── Base ─────────────────────────────────────────────────── */
QMainWindow {{
    background-color: {bg};
    border: 1px solid {brd};
}}
QWidget {{
    color: {txt};
    font-family: {font_family};
    font-size: 13px;
    background-color: transparent;
}}
QWidget#root {{
    background-color: {bg};
}}

/* ── Title Bar ────────────────────────────────────────── */
QFrame#custom-titlebar {{
    background-color: {card};
    border-bottom: 1px solid {brd};
}}
QPushButton#btn-titlebar {{
    background-color: transparent;
    border: none;
    font-size: 14px;
    color: {txt};
}}
QPushButton#btn-titlebar:hover {{
    background-color: {subtle};
}}
QPushButton#btn-titlebar-close {{
    background-color: transparent;
    border: none;
    font-size: 14px;
    color: {txt};
}}
QPushButton#btn-titlebar-close:hover {{
    background-color: #e81123;
    color: white;
}}

/* ── Cards ────────────────────────────────────────────────── */
QFrame#card {{
    background-color: {card};
    border: 1px solid {brd};
    border-radius: 8px;
}}

/* ── Labels ───────────────────────────────────────────────── */
QLabel#section-header {{
    color: {muted};
    font-size: 10px;
    font-weight: 600;
    letter-spacing: 1.2px;
}}
QLabel#sublabel {{
    color: {muted};
    font-size: 11px;
}}
QLabel#status-label {{
    color: {a};
    font-weight: 600;
    font-size: 12px;
}}
QLabel#timer-label {{
    color: {dim};
    font-size: 12px;
    font-family: 'Cascadia Code', 'Consolas', monospace;
}}
QLabel#footer-text {{
    color: {dim};
    font-size: 11px;
}}
QLabel#link {{
    color: {a};
    font-size: 11px;
}}

/* ── Inputs ───────────────────── */
QLineEdit, QComboBox, QSpinBox {{
    background-color: {inp};
    border: 1px solid {brd};
    border-radius: 6px;
    padding: 5px 10px;
    color: {txt};
    min-height: 28px;
    selection-background-color: {a};
    selection-color: {bg};
}}
QLineEdit:focus, QComboBox:focus, QSpinBox:focus {{
    border-color: {a};
}}
QLineEdit:read-only {{
    color: {muted};
    background-color: {subtle};
}}
QLineEdit:disabled, QComboBox:disabled, QSpinBox:disabled {{
    color: {inp_dis_txt};
    border-color: {inp_dis_brd};
    background-color: {subtle};
}}

/* ── Pop-ups ────────────────── */
QMessageBox {{
    background-color: {bg};
}}
QMessageBox QLabel {{
    color: {txt};
    background-color: transparent;
}}
QMessageBox QPushButton {{
    background-color: {subtle};
    color: {txt};
    border: 1px solid {brd};
    border-radius: 6px;
    padding: 6px 16px;
    min-height: 26px;
    font-weight: 600;
}}
QMessageBox QPushButton:hover {{
    background-color: {brd};
}}

/* ── ComboBox / Dropdown ──────────────────────────────────── */
QComboBox::drop-down {{
    subcontrol-origin: padding;
    subcontrol-position: top right;
    width: 22px;
    border: none;
    background: transparent;
}}
QComboBox::down-arrow {{
    image: url("data:image/svg+xml;utf8,<svg viewBox='0 0 24 24' xmlns='http://www.w3.org/2000/svg'><path d='M7 10l5 5 5-5z' fill='{muted_svg}'/></svg>");
    width: 14px; 
    height: 14px; 
    margin-right: 6px;
}}
QComboBox QAbstractItemView {{
    background-color: {card};
    border: 1px solid {brd};
    border-radius: 6px;
    selection-background-color: {subtle};
    selection-color: {a};
    outline: none;
    padding: 2px;
}}

/* ── SpinBox ──────────────────────────────────────────────── */
QSpinBox::up-button, QSpinBox::down-button {{
    border: none;
    background: transparent;
    width: 16px;
}}
QSpinBox::up-arrow {{
    image: url("data:image/svg+xml;utf8,<svg viewBox='0 0 24 24' xmlns='http://www.w3.org/2000/svg'><path d='M7 14l5-5 5 5z' fill='{muted_svg}'/></svg>");
    width: 12px; 
    height: 12px;
}}
QSpinBox::down-arrow {{
    image: url("data:image/svg+xml;utf8,<svg viewBox='0 0 24 24' xmlns='http://www.w3.org/2000/svg'><path d='M7 10l5 5 5-5z' fill='{muted_svg}'/></svg>");
    width: 12px; 
    height: 12px;
}}

/* ── Checkboxes ───────────────────────────────────────────── */
QCheckBox {{
    spacing: 8px;
    color: {txt};
}}
QCheckBox::indicator {{
    width: 15px;
    height: 15px;
    border: 1px solid {brd};
    border-radius: 4px;
    background: {inp};
}}
QCheckBox::indicator:checked {{
    background-color: {a};
    border-color: {a};
}}
QCheckBox::indicator:hover {{
    border-color: {muted};
}}
QCheckBox:disabled {{
    color: {dim};
}}

/* ── Primary button ───────────────────────────────────────── */
QPushButton#btn-primary {{
    background-color: {a};
    color: {bg if dark else "#ffffff"};
    font-weight: 700;
    border: none;
    border-radius: 6px;
    padding: 0px 16px;
    min-height: 38px;
    font-size: 13px;
}}
QPushButton#btn-primary:hover   {{ background-color: {a_hover}; }}
QPushButton#btn-primary:pressed {{ background-color: {a_press}; }}
QPushButton#btn-primary:disabled {{
    background-color: {a_dis_bg};
    color: {a_dis_fg};
}}

/* ── Danger button ────────────────────────────────────────── */
QPushButton#btn-danger {{
    background-color: transparent;
    color: {danger_fg};
    font-weight: 600;
    border: 1px solid {danger_brd};
    border-radius: 6px;
    padding: 0px 16px;
    min-height: 38px;
    font-size: 13px;
}}
QPushButton#btn-danger:hover {{
    background-color: {danger_hover};
    border-color: {danger_hover};
    color: #ffffff;
}}
QPushButton#btn-danger:disabled {{
    color: {danger_dis_fg};
    border-color: {danger_dis_brd};
}}

/* ── Secondary button ─────────────────────────────────────── */
QPushButton#btn-secondary {{
    background-color: {subtle};
    color: {txt};
    border: 1px solid {brd};
    border-radius: 6px;
    padding: 0px 12px;
    min-height: 30px;
}}
QPushButton#btn-secondary:hover   {{ background-color: {brd}; }}
QPushButton#btn-secondary:pressed {{ background-color: {muted}; color: {bg}; }}

/* ── Header controls ──────────────────────────────────────── */
QPushButton#btn-header {{
    background-color: transparent;
    color: {muted};
    border: 1px solid {brd_sub};
    border-radius: 4px;
    padding: 0px 8px;
    min-height: 24px;
    font-size: 11px;
    font-weight: 600;

}}
QPushButton#btn-header:hover {{
    border-color: {brd};
    color: {txt};
    background-color: {subtle};
}}

/* ── Progress bar ─────────────────────────────────────────── */
QProgressBar {{
    background-color: {subtle};
    border: none;
    border-radius: 3px;
    max-height: 5px;
    text-align: center;
    color: transparent;
}}
QProgressBar::chunk {{
    background-color: {a};
    border-radius: 3px;
}}

/* ── Log ──────────────────────────────────────────────────── */
QPlainTextEdit#log {{
    background-color: {log_bg};
    border: none;
    color: {log_txt};
    font-family: 'Cascadia Code', 'Fira Code', 'Consolas', monospace;
    font-size: 12px;
    padding: 12px;
    selection-background-color: {a};
    selection-color: {bg};
}}

/* ── Scrollbars ───────────────────────────────────────────── */
QScrollBar:vertical {{
    background: {log_bg};
    width: 7px;
    margin: 0;
}}
QScrollBar::handle:vertical {{
    background: {brd};
    border-radius: 3px;
    min-height: 24px;
}}

QScrollBar::handle:vertical:hover  {{ background: {muted}; }}

QScrollBar::add-line:vertical, 
QScrollBar::sub-line:vertical {{ height: 0; }}

QScrollBar::add-page:vertical, 
QScrollBar::sub-page:vertical {{ background: none; }}

/* ── Frame containers ─────────────────────────────────────── */
QFrame#header {{
    background-color: {card};
    border-bottom: 1px solid {brd};
}}
QFrame#footer {{
    background-color: {bg};
    border-top: 1px solid {brd_sub};
}}
QFrame#log-header {{
    background-color: {card};
    border-bottom: 1px solid {brd_sub};
}}
QFrame#theme-swatch {{
    border-radius: 7px;
    border: 2px solid transparent;
}}
QFrame#theme-swatch:hover {{
    border-color: {txt};
}}
"""

class CustomTitleBar(QFrame):
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        self.setObjectName("custom-titlebar")
        self.setFixedHeight(36)
        
        layout = QHBoxLayout(self)
        layout.setContentsMargins(15, 0, 0, 0)
        layout.setSpacing(0)
        
        # Başlık Metni
        self.title_label = QLabel("Whisper Transcript")
        self.title_label.setStyleSheet("font-weight: 700; font-size: 13px;")
        layout.addWidget(self.title_label)
        
        layout.addStretch()
        
        # Simge Durumuna Küçült Butonu
        self.btn_min = QPushButton("─")
        self.btn_min.setObjectName("btn-titlebar")
        self.btn_min.setFixedSize(46, 36)
        self.btn_min.clicked.connect(self.parent.showMinimized)
        layout.addWidget(self.btn_min)
        
        # Büyült/Küçült Butonu
        self.btn_max = QPushButton("◻")
        self.btn_max.setObjectName("btn-titlebar")
        self.btn_max.setFixedSize(46, 36)
        self.btn_max.clicked.connect(self.toggle_maximize)
        layout.addWidget(self.btn_max)
        
        # Kapat Butonu
        self.btn_close = QPushButton("✕")
        self.btn_close.setObjectName("btn-titlebar-close")
        self.btn_close.setFixedSize(46, 36)
        self.btn_close.clicked.connect(self.parent.close)
        layout.addWidget(self.btn_close)
        
        self.start_pos = None

    def toggle_maximize(self):
        if self.parent.isMaximized():
            self.parent.showNormal()
        else:
            self.parent.showMaximized()

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.start_pos = event.globalPosition().toPoint()

    def mouseMoveEvent(self, event):
        if self.start_pos:
            # Sürükleme esnasında pencere maksimize ise önce normale döndür
            if self.parent.isMaximized():
                self.parent.showNormal()
                self.start_pos = QPoint(int(self.parent.width() / 2), 10)
                
            delta = event.globalPosition().toPoint() - self.start_pos
            self.parent.move(self.parent.pos() + delta)
            self.start_pos = event.globalPosition().toPoint()

    def mouseReleaseEvent(self, event):
        self.start_pos = None

    def mouseDoubleClickEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.toggle_maximize()

# ── Stderr capture for standard Whisper tqdm progress ────────────────────────
class TqdmCapture:
    def __init__(self, progress_callback):
        self._callback = progress_callback
        self._buffer   = ""

    def write(self, text):
        self._buffer += text
        match = re.search(r"(\d+)/(\d+)", self._buffer)
        if match:
            current = int(match.group(1))
            total   = int(match.group(2))
            if total > 0:
                self._callback(current / total * 100)
            self._buffer = self._buffer[-20:]

    def flush(self):
        pass


# ── Update checker thread ─────────────────────────────────────────────────────
class UpdateChecker(QThread):
    update_available = pyqtSignal(str, object)  # new_version_str, asset_dict_or_None

    def run(self):
        try:
            api_url = f"https://api.github.com/repos/{GITHUB_REPO}/releases/latest"
            req = urllib.request.Request(
                api_url,
                headers={
                    "User-Agent": "whisper-gui-updater",
                    "Accept":     "application/vnd.github+json",
                }
            )
            with urllib.request.urlopen(req, timeout=8) as resp:
                data = json.loads(resp.read().decode("utf-8"))

            tag = data.get("tag_name", "").strip().lstrip("v")
            if not tag:
                return

            current = tuple(int(x) for x in APP_VERSION.split("."))
            latest  = tuple(int(x) for x in tag.split("."))
            length  = max(len(current), len(latest))
            current = current + (0,) * (length - len(current))
            latest  = latest  + (0,) * (length - len(latest))

            if latest <= current:
                return

            assets    = data.get("assets", [])
            pyw_asset = next(
                (a for a in assets if a.get("name", "").endswith(".pyw")), None
            )
            self.update_available.emit(tag, pyw_asset)

        except Exception:
            pass


# ── Transcript worker thread ──────────────────────────────────────────────────
class TranscriptWorker(QThread):
    log      = pyqtSignal(str)
    progress = pyqtSignal(float)
    status   = pyqtSignal(str)
    done     = pyqtSignal(bool, float)
    error    = pyqtSignal(str, str)   # title, message

    def __init__(self, params, model_ref, model_lock):
        super().__init__()
        self._params      = params
        self._model_ref   = model_ref   # {"model": obj, "key": str}
        self._model_lock  = model_lock
        self._cancel_flag = False
        self._temp_path   = ""

    def request_cancel(self):
        self._cancel_flag = True

    # ── Main run ──────────────────────────────────────────────────────────────
    def run(self):
        original_stderr = sys.stderr
        start_time      = time.time()
        params          = self._params

        try:
            # ── Install mode ──────────────────────────────────────────────
            if params["install_mode"]:
                self.log.emit(">>> INSTALLATION STARTED...")
                for lib in LIBS_TO_INSTALL:
                    self.log.emit(f"   -> Installing: {lib}")
                    try:
                        subprocess.run(
                            [sys.executable, "-m", "pip", "install", lib],
                            capture_output=True, text=True, check=True,
                        )
                    except subprocess.CalledProcessError as e:
                        raise Exception(
                            f"{lib} installation failed.\nDetail: {e.stderr}"
                        )

                self.log.emit("   -> Installing: PyTorch (CUDA)")
                try:
                    subprocess.run(
                        [sys.executable, "-m", "pip", "install"] + PYTORCH_INSTALL_ARGS,
                        capture_output=True, text=True, check=True,
                    )
                except subprocess.CalledProcessError:
                    raise Exception("PyTorch (CUDA) installation failed.")

                self.log.emit(">>> Installation complete.\n")

            # ── Resolve model ─────────────────────────────────────────────
            model_display = params["model_display"]
            model_key     = params["model_map"].get(model_display)
            if not model_key:
                raise ValueError(f"Invalid model selection: '{model_display}'")

            target_path = MODEL_PATHS[model_key]

            # ── Download model if missing ─────────────────────────────────
            if not os.path.exists(target_path):
                self.log.emit(f"\nDOWNLOADING: {target_path}")
                self.status.emit("Downloading model...")
                if "faster" in model_key:
                    from faster_whisper import download_model
                    repo_id = "turbo" if "turbo" in model_key else "large-v3"
                    download_model(repo_id, output_dir=target_path)
                    self.log.emit("Download complete!")
                else:
                    import whisper
                    model_name = model_key.replace("std_", "").replace("large", "large-v3")
                    url = whisper._MODELS[model_name]
                    from tqdm import tqdm
                    with tqdm(unit="B", unit_scale=True, miniters=1, desc=model_name) as t:
                        def _hook(blocknum, blocksize, totalsize):
                            t.total = totalsize
                            t.update(blocknum * blocksize - t.n)
                        urllib.request.urlretrieve(url, target_path, _hook)
                    self.log.emit("Download complete!")

            # ── Load model ────────────────────────────────────────────────
            with self._model_lock:
                if self._model_ref.get("key") != model_key:
                    existing = self._model_ref.get("model")
                    if existing is not None:
                        self.log.emit(">>> Unloading previous model from VRAM...")
                        del existing
                        self._model_ref["model"] = None
                        self._model_ref["key"]   = None
                        gc.collect()
                        try:
                            import torch
                            if torch.cuda.is_available():
                                torch.cuda.empty_cache()
                        except ImportError:
                            pass
                    self._model_ref["key"] = model_key
                    need_load = True
                else:
                    need_load = False
                    self.log.emit(">>> Model already in VRAM, reusing...")

            if "faster" in model_key:
                from faster_whisper import WhisperModel
                if need_load:
                    self.status.emit("Loading model...")
                    loaded = WhisperModel(
                        target_path,
                        device="cuda",
                        compute_type=params["compute_type"],
                        local_files_only=True,
                    )
                    with self._model_lock:
                        self._model_ref["model"] = loaded
            else:
                import whisper
                if need_load:
                    self.status.emit("Loading model...")
                    load_arg = (
                        target_path if os.path.exists(target_path)
                        else model_key.replace("std_", "").replace("large", "large-v3")
                    )
                    loaded = whisper.load_model(load_arg, device="cuda")
                    with self._model_lock:
                        self._model_ref["model"] = loaded

            model_obj   = self._model_ref["model"]
            files       = params["files"]
            total_files = len(files)

            # ── Process each file ─────────────────────────────────────────
            for idx, file_path in enumerate(files):
                if self._cancel_flag:
                    break

                self.progress.emit(0)
                self.status.emit(
                    f"Processing {idx+1}/{total_files}..."
                    if total_files > 1 else "Processing..."
                )

                output_path = (
                    params["output_path"]
                    if total_files == 1
                    else self._build_output_path(file_path, model_key, params)
                )

                temp_path       = os.path.splitext(output_path)[0] + "_temp.txt"
                self._temp_path = temp_path
                with open(temp_path, "w", encoding="utf-8"):
                    pass

                # Use index-based language lookup so UI language doesn't matter
                lang_code    = LANGUAGE_CODES[params["language_index"]]
                out_format   = params["output_format"]
                vad_active   = params["vad"]
                timestamps   = params["timestamps"]
                single_block = params["single_block"]
                translate    = params["translate"]
                task         = "translate" if translate else "transcribe"

                if "XML" in out_format and not timestamps:
                    timestamps = True

                self.log.emit(f"\n{'='*50}")
                self.log.emit(f"[{idx+1}/{total_files}] {os.path.basename(file_path)}")
                self.log.emit(
                    f"Model: {model_key} | Compute: {params['compute_type']}"
                    f" | Task: {task.upper()}"
                )
                self.log.emit(f"{'='*50}\n")

                file_start = time.time()

                if "faster" in model_key:
                    transcribe_args = {
                        "beam_size":                  params["beam_size"],
                        "language":                   lang_code,
                        "task":                       task,
                        "condition_on_previous_text": not vad_active,
                        "vad_filter":                 vad_active,
                        "word_timestamps":            False,
                    }
                    if vad_active:
                        transcribe_args["vad_parameters"] = dict(
                            min_silence_duration_ms=500
                        )
                        try:
                            transcribe_args["repetition_penalty"] = 1.1
                        except Exception as rp_err:
                            self.log.emit(
                                f"! Warning: repetition_penalty not supported: {rp_err}"
                            )

                    segments, info = model_obj.transcribe(file_path, **transcribe_args)
                    self.log.emit(f"Detected language: {info.language.upper()}\n")

                    with open(temp_path, "w", encoding="utf-8", buffering=1) as f:
                        for seg in segments:
                            if info.duration > 0:
                                self.progress.emit(seg.end / info.duration * 100)
                            text = seg.text.strip()
                            if timestamps:
                                line = f"{self._fmt_timestamp(seg.start)} {text}\n"
                            elif single_block:
                                line = f"{text} "
                            else:
                                line = f"{text}\n"
                            f.write(line)
                            self.log.emit(line.strip())
                            if self._cancel_flag:
                                break

                else:  # Standard Whisper
                    capture     = TqdmCapture(lambda p: self.progress.emit(p))
                    sys.stderr  = capture
                    try:
                        result = model_obj.transcribe(
                            file_path,
                            language=lang_code,
                            task=task,
                            verbose=False,
                            condition_on_previous_text=not vad_active,
                            no_speech_threshold=0.6,
                            beam_size=params["beam_size"],
                        )
                        with open(temp_path, "w", encoding="utf-8", buffering=1) as f:
                            for seg in result["segments"]:
                                text = seg["text"].strip()
                                if timestamps:
                                    line = f"{self._fmt_timestamp(seg['start'])} {text}\n"
                                elif single_block:
                                    line = f"{text} "
                                else:
                                    line = f"{text}\n"
                                f.write(line)
                                self.log.emit(
                                    text[:100] + "..." if single_block else line.strip()
                                )
                    finally:
                        sys.stderr = original_stderr

                # ── Save output ───────────────────────────────────────────
                if "TXT" in out_format:
                    if os.path.exists(output_path):
                        os.remove(output_path)
                    os.rename(temp_path, output_path)
                else:
                    if self._convert_file(temp_path, output_path, out_format):
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                    else:
                        output_path = temp_path

                self._temp_path = ""
                elapsed = time.time() - file_start
                self.log.emit(f"\n✓ File time: {elapsed:.2f}s")

            # ── Finish ────────────────────────────────────────────────────
            total_elapsed = time.time() - start_time
            if self._cancel_flag:
                self.status.emit("Stopped")
                self.log.emit("\nProcess stopped by user.")
                self.done.emit(False, total_elapsed)
            else:
                self.log.emit(
                    f"\n{'='*50}\nTOTAL TIME: {total_elapsed:.2f}s\n{'='*50}"
                )
                self.done.emit(True, total_elapsed)

        except RuntimeError as e:
            msg = str(e).lower()
            if any(k in msg for k in ("memory", "cuda out of memory", "allocate")):
                with self._model_lock:
                    m = self._model_ref.get("model")
                    if m:
                        del m
                        self._model_ref["model"] = None
                        self._model_ref["key"]   = None
                gc.collect()
                try:
                    import torch
                    if torch.cuda.is_available():
                        torch.cuda.empty_cache()
                except ImportError:
                    pass
                self.error.emit(
                    "Out of Memory",
                    f"VRAM exhausted.\nTry switching Compute to 'int8'.\n\nDetail:\n{e}",
                )
            else:
                self._handle_error(e)

        except Exception as e:
            self._handle_error(e)

        finally:
            sys.stderr = original_stderr
            if self._temp_path and os.path.exists(self._temp_path):
                try:
                    os.remove(self._temp_path)
                    self.log.emit(
                        f"! Temp file cleaned: {os.path.basename(self._temp_path)}"
                    )
                except OSError:
                    pass
            self._temp_path = ""

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _handle_error(self, e):
        msg = str(e).lower()
        if any(k in msg for k in ("no cuda-capable device", "cuda failed", "cublas")):
            self.error.emit(
                "GPU Not Found",
                "No NVIDIA GPU was detected or the driver is inactive.\n"
                "Please ensure you have a CUDA-capable GPU with up-to-date drivers.",
            )
        else:
            self.error.emit("Process Error", f"Unexpected error:\n\n{str(e)}")

    def _fmt_timestamp(self, seconds):
        h  = int(seconds // 3600)
        m  = int((seconds % 3600) // 60)
        s  = int(seconds % 60)
        return f"[{h:02}:{m:02}:{s:02}]"

    def _build_output_path(self, input_file, model_key, params):
        folder, name = os.path.split(input_file)
        stem         = os.path.splitext(name)[0]
        fmt          = params["output_format"]
        ext_map      = {"PDF": ".pdf", "DOCX": ".docx", "XML": ".xml"}
        ext          = next((v for k, v in ext_map.items() if k in fmt), ".txt")
        trans_suffix = "_translated" if params["translate"] else ""
        return os.path.join(
            folder,
            f"{stem}_{model_key.replace('_', '-')}{trans_suffix}_transcript{ext}",
        )

    def _convert_file(self, txt_path, target_path, format_type):
        try:
            def ensure_lib(package, import_name):
                try:
                    __import__(import_name)
                    return True
                except ImportError:
                    self.log.emit(f">>> Installing missing library: {package}...")
                    try:
                        si = subprocess.STARTUPINFO()
                        si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                        subprocess.run(
                            [sys.executable, "-m", "pip", "install", package],
                            check=True, capture_output=True, startupinfo=si,
                        )
                        __import__(import_name)
                        self.log.emit(f"   -> {package} installed.")
                        return True
                    except subprocess.CalledProcessError:
                        self.log.emit(f"!!! Failed to install {package}.")
                        return False

            with open(txt_path, "r", encoding="utf-8") as f:
                content = f.read()

            if "DOCX" in format_type:
                if not ensure_lib("python-docx", "docx"):
                    raise Exception("python-docx could not be installed.")
                self.log.emit(">>> Creating Word (DOCX)...")
                from docx import Document
                doc = Document()
                doc.add_heading("Transcript", 0)
                for line in content.split("\n"):
                    stripped = line.strip()
                    if stripped:
                        doc.add_paragraph(stripped)
                doc.save(target_path)

            elif "PDF" in format_type:
                if not ensure_lib("fpdf", "fpdf"):
                    raise Exception("fpdf could not be installed.")
                self.log.emit(">>> Creating PDF...")
                from fpdf import FPDF
                font_candidates = [
                    r"C:\Windows\Fonts\arial.ttf",
                    r"/Library/Fonts/Arial.ttf",
                    r"/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",
                    r"/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                ]
                font_path = next(
                    (p for p in font_candidates if os.path.exists(p)), None
                )

                class _PDF(FPDF):
                    def header(self):
                        if font_path:
                            try:
                                self.add_font("CustomFont", "", font_path, uni=True)
                                self.set_font("CustomFont", "", 10)
                            except Exception:
                                self.set_font("Arial", "", 10)
                        else:
                            self.set_font("Arial", "", 10)
                        self.cell(0, 10, "Transcript", 0, 1, "C")

                pdf = _PDF()
                pdf.add_page()
                if font_path:
                    try:
                        pdf.add_font("CustomFont", "", font_path, uni=True)
                        pdf.set_font("CustomFont", size=11)
                    except Exception as fe:
                        self.log.emit(f"! Font warning: {fe}")
                        pdf.set_font("Arial", size=11)
                        content = content.encode("latin-1", "replace").decode("latin-1")
                else:
                    self.log.emit("! No UTF-8 font found; characters may be garbled.")
                    pdf.set_font("Arial", size=11)
                    content = content.encode("latin-1", "replace").decode("latin-1")
                pdf.multi_cell(0, 8, content)
                pdf.output(target_path)

            elif "XML" in format_type:
                self.log.emit(">>> Creating XML (Subtitle)...")
                xml_content = self._build_xml(content)
                with open(target_path, "w", encoding="utf-8") as f:
                    f.write(xml_content)

            self.log.emit(f"✓ Converted: {os.path.basename(target_path)}")
            return True

        except Exception as e:
            self.log.emit(f"!!! Conversion error: {e}")
            return False

    def _build_xml(self, text):
        lines   = [
            '<?xml version="1.0" encoding="UTF-8"?>',
            '<tt xmlns="http://www.w3.org/ns/ttml" xml:lang="tr">',
            "  <body>", "    <div>",
        ]
        pattern = re.compile(r"\[(\d{2}):(\d{2}):(\d{2})\]\s*(.*)")
        parsed  = []
        for raw in text.strip().split("\n"):
            m = pattern.match(raw.strip())
            if m:
                h, mi, s = int(m.group(1)), int(m.group(2)), int(m.group(3))
                parsed.append({"time": h * 3600 + mi * 60 + s, "text": m.group(4)})

        def _fmt(sec):
            h  = int(sec // 3600)
            mi = int((sec % 3600) // 60)
            s  = int(sec % 60)
            ms = int((sec - int(sec)) * 1000)
            return f"{h:02}:{mi:02}:{s:02}.{ms:03}"

        for i, entry in enumerate(parsed):
            end_sec = (
                parsed[i + 1]["time"] if i < len(parsed) - 1 else entry["time"] + 3.0
            )
            safe = (
                entry["text"]
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            lines.append(
                f'       <p begin="{_fmt(entry["time"])}" end="{_fmt(end_sec)}">{safe}</p>'
            )

        lines.extend(["    </div>", "  </body>", "</tt>"])
        return "\n".join(lines)


# ── Main window ───────────────────────────────────────────────────────────────
class WhisperApp(QMainWindow):
    def __init__(self):
        super().__init__()
        
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Window)

        # Shared model state
        self._model_ref  = {"model": None, "key": None}
        self._model_lock = threading.Lock()
        self._worker     = None

        # UI state
        self._selected_files = []
        self._model_map      = {}
        self._timer_active   = False
        self._timer_start    = 0.0

        # Appearance state
        self._ui_lang    = "tr"
        self._theme_name = "blue"
        self._dark_mode  = True

        self._load_preferences()

        self._check_ffmpeg()
        self._build_ui()
        self._btn_mode.setText("☾" if self._dark_mode else "☀")
        self._apply_theme(animate=False)
        self._scan_models()
        self._setup_timer()

        updater = UpdateChecker(self)
        updater.update_available.connect(self._on_update_available)
        updater.start()

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _s(self, key):
        """Return the current-language string for the given key."""
        return STRINGS[self._ui_lang][key]

    # ── FFmpeg check ──────────────────────────────────────────────────────────
    def _check_ffmpeg(self):
        if shutil.which("ffmpeg"):
            return
        reply = QMessageBox.question(
            self, "FFmpeg Missing",
            "FFmpeg was not found on your system. "
            "It is required for audio processing.\n\n"
            "Download and install it automatically from GitHub?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if reply != QMessageBox.StandardButton.Yes:
            QMessageBox.warning(self, "Warning",
                "FFmpeg is missing. The application may not work correctly.")
            return
        try:
            script = "ffmpeg_installer_temp.bat"
            urllib.request.urlretrieve(FFMPEG_BAT_URL, script)
            subprocess.check_call([script], shell=True)
            if os.path.exists(script):
                os.remove(script)
            QMessageBox.information(self, "Restarting",
                "FFmpeg installed. The application will restart.")
            subprocess.Popen([sys.executable] + sys.argv)
            sys.exit(0)
        except Exception as e:
            QMessageBox.critical(self, "Install Error",
                f"Automatic installation failed:\n{e}")

    # ── UI construction ───────────────────────────────────────────────────────
    def _build_ui(self):
        self.setWindowTitle("Whisper Transcript")
        self.setMinimumSize(1100, 750)
        self.resize(1280, 820)

        root = QWidget()
        root.setObjectName("root")
        self.setCentralWidget(root)

        layout = QVBoxLayout(root)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        self.titlebar = CustomTitleBar(self)
        layout.addWidget(self.titlebar)

        layout.addWidget(self._make_header())

        body = QWidget()
        body.setObjectName("root")
        body_layout = QHBoxLayout(body)
        body_layout.setContentsMargins(16, 16, 16, 10)
        body_layout.setSpacing(12)

        sidebar = self._make_sidebar()
        sidebar.setFixedWidth(390)
        body_layout.addWidget(sidebar)
        body_layout.addWidget(self._make_log_panel(), stretch=1)

        layout.addWidget(body, stretch=1)
        layout.addWidget(self._make_footer())
        
        size_grip_layout = QHBoxLayout()
        size_grip_layout.setContentsMargins(0, 0, 0, 0)
        size_grip_layout.addStretch()
        size_grip = QSizeGrip(self)
        size_grip_layout.addWidget(size_grip)
        layout.addLayout(size_grip_layout)

    # ── Header ────────────────────────────────────────────────────────────────
    def _make_header(self):
        header = QFrame()
        header.setObjectName("header")
        header.setFixedHeight(56)
        hl = QHBoxLayout(header)
        hl.setContentsMargins(20, 0, 16, 0)
        hl.setSpacing(6)

        self._lbl_title = QLabel("Whisper")
        self._lbl_title.setStyleSheet(
            "font-size:19px; font-weight:700; letter-spacing:-0.5px;"
        )
        self._lbl_subtitle = QLabel(self._s("header_sub"))
        self._lbl_subtitle.setStyleSheet("font-size:14px;")

        hl.addWidget(self._lbl_title)
        hl.addWidget(self._lbl_subtitle)
        hl.addStretch()

        # ── Theme color swatches ──────────────────────────────────────────
        swatch_colors = {
            "yellow": "#f0a500",
            "red":    "#ef4444",
            "green":  "#22c55e",
            "blue":   "#3b82f6",
        }
        self._swatch_btns = {}
        for name, color in swatch_colors.items():
            btn = QPushButton()
            btn.setFixedSize(14, 14)
            btn.setStyleSheet(
                f"QPushButton {{ background-color: {color}; border-radius: 7px; border: none; }}"
                f"QPushButton:hover {{ border: 2px solid white; }}"
            )
            btn.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
            btn.clicked.connect(lambda _, n=name: self._set_theme_color(n))
            self._swatch_btns[name] = btn
            hl.addWidget(btn)

        hl.addSpacing(10)

        # ── Dark / Light toggle ───────────────────────────────────────────
        self._btn_mode = QPushButton("☾")
        self._btn_mode.setObjectName("btn-header")
        self._btn_mode.setFixedSize(32, 24)
        self._btn_mode.setToolTip("Toggle dark / light mode")
        self._btn_mode.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self._btn_mode.clicked.connect(self._toggle_mode)
        hl.addWidget(self._btn_mode)

        hl.addSpacing(4)

        # ── Dynamic Language Dropdown ─────────────────
        self._lang_ui_combo = QComboBox()
        self._lang_ui_combo.setFixedWidth(90)
        self._lang_ui_combo.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        
        # Integrate the new languages
        self._ui_langs_map = {"Türkçe": "tr", "English": "en", "Deutsch": "de", "Français": "fr", "Español": "es", "中文": "zh", "日本語": "ja"}
        self._lang_ui_combo.addItems(list(self._ui_langs_map.keys()))
        
        # Apply current language automatically
        current_name = [k for k, v in self._ui_langs_map.items() if v == self._ui_lang]
        if current_name:
            self._lang_ui_combo.setCurrentText(current_name[0])
            
        self._lang_ui_combo.currentTextChanged.connect(self._change_language)
        hl.addWidget(self._lang_ui_combo)

        hl.addSpacing(8)

        self._lbl_ver = QLabel(f"v{APP_VERSION}")
        self._lbl_ver.setStyleSheet(
            "font-size:11px; font-family:monospace;"
        )
        hl.addWidget(self._lbl_ver)

        return header

    # ── Sidebar ───────────────────────────────────────────────────────────────
    def _make_sidebar(self):
        w  = QWidget()
        w.setObjectName("root")
        vl = QVBoxLayout(w)
        vl.setContentsMargins(0, 0, 0, 0)
        vl.setSpacing(10)

        vl.addWidget(self._make_file_card())
        vl.addWidget(self._make_settings_card())
        vl.addWidget(self._make_options_card())
        vl.addWidget(self._make_progress_card())
        vl.addWidget(self._make_controls())
        vl.addStretch()
        return w

    # ── Card factory ──────────────────────────────────────────────────────────
    def _make_card(self, text_key):
        if not hasattr(self, "_section_labels"):
            self._section_labels = {}
        card   = QFrame()
        card.setObjectName("card")
        layout = QVBoxLayout(card)
        layout.setContentsMargins(14, 10, 14, 14)
        layout.setSpacing(8)

        lbl = QLabel(self._s(text_key).upper())
        lbl.setObjectName("section-header")
        layout.addWidget(lbl)

        self._section_labels[text_key] = lbl
        return card, layout

    # ── File & format card ────────────────────────────────────────────────────
    def _make_file_card(self):
        card, vl = self._make_card("sec_files")

        src_row = QHBoxLayout()
        self._file_display = QLineEdit()
        self._file_display.setPlaceholderText(self._s("ph_files"))
        self._file_display.setReadOnly(True)
        self._btn_browse = QPushButton(self._s("btn_browse"))
        self._btn_browse.setObjectName("btn-secondary")
        self._btn_browse.setFixedWidth(72)
        self._btn_browse.clicked.connect(self._select_files)
        src_row.addWidget(self._file_display)
        src_row.addWidget(self._btn_browse)
        vl.addLayout(src_row)

        fmt_row = QHBoxLayout()
        self._format_combo = QComboBox()
        self._format_combo.addItems(FORMAT_OPTIONS)
        self._format_combo.setFixedWidth(128)
        self._format_combo.currentTextChanged.connect(self._on_format_change)
        fmt_row.addWidget(self._format_combo)

        self._output_display = QLineEdit()
        self._output_display.setPlaceholderText(self._s("ph_output"))
        self._btn_out = QPushButton("…")
        self._btn_out.setObjectName("btn-secondary")
        self._btn_out.setFixedWidth(30)
        self._btn_out.clicked.connect(self._select_output)
        fmt_row.addWidget(self._output_display)
        fmt_row.addWidget(self._btn_out)
        vl.addLayout(fmt_row)

        return card

    # ── Model & settings card ─────────────────────────────────────────────────
    def _make_settings_card(self):
        card, vl = self._make_card("sec_model")

        self._model_combo = QComboBox()
        self._model_combo.currentIndexChanged.connect(self._on_model_change)
        vl.addWidget(self._model_combo)

        row = QHBoxLayout()
        row.setSpacing(10)

        lang_col = QVBoxLayout()
        self._lbl_lang_field = QLabel(self._s("lbl_language"))
        self._lbl_lang_field.setObjectName("sublabel")
        self._lang_combo = QComboBox()
        self._lang_combo.addItems(self._s("lang_names"))
        lang_col.addWidget(self._lbl_lang_field)
        lang_col.addWidget(self._lang_combo)
        row.addLayout(lang_col, stretch=2)

        compute_col = QVBoxLayout()
        self._lbl_compute = QLabel(self._s("lbl_compute"))
        self._lbl_compute.setObjectName("sublabel")
        self._compute_combo = QComboBox()
        self._compute_combo.addItems(COMPUTE_OPTIONS)
        compute_col.addWidget(self._lbl_compute)
        compute_col.addWidget(self._compute_combo)
        row.addLayout(compute_col, stretch=2)

        beam_col = QVBoxLayout()
        self._lbl_beam = QLabel(self._s("lbl_beam"))
        self._lbl_beam.setObjectName("sublabel")
        self._beam_spin = QSpinBox()
        self._beam_spin.setRange(1, 10)
        self._beam_spin.setValue(5)
        beam_col.addWidget(self._lbl_beam)
        beam_col.addWidget(self._beam_spin)
        row.addLayout(beam_col, stretch=1)

        vl.addLayout(row)
        return card

    # ── Options card ──────────────────────────────────────────────────────────
    def _make_options_card(self):
        card, vl = self._make_card("sec_options")

        grid = QGridLayout()
        grid.setHorizontalSpacing(16)
        grid.setVerticalSpacing(8)

        self._chk_vad        = QCheckBox(self._s("chk_vad"))
        self._chk_single     = QCheckBox(self._s("chk_single"))
        self._chk_timestamps = QCheckBox(self._s("chk_timestamps"))
        self._chk_translate  = QCheckBox(self._s("chk_translate"))
        self._chk_install    = QCheckBox(self._s("chk_install"))

        self._chk_vad.setChecked(True)
        self._chk_single.setChecked(True)

        self._chk_single.toggled.connect(
            lambda on: self._chk_timestamps.setChecked(False) if on else None
        )
        self._chk_timestamps.toggled.connect(
            lambda on: self._chk_single.setChecked(False) if on else None
        )

        grid.addWidget(self._chk_vad,        0, 0)
        grid.addWidget(self._chk_single,     0, 1)
        grid.addWidget(self._chk_timestamps, 1, 0)
        grid.addWidget(self._chk_translate,  1, 1)
        grid.addWidget(self._chk_install,    2, 0)

        vl.addLayout(grid)
        return card

    # ── Progress card ─────────────────────────────────────────────────────────
    def _make_progress_card(self):
        card, vl = self._make_card("sec_progress")

        status_row = QHBoxLayout()
        self._status_label = QLabel(self._s("status_ready"))
        self._status_label.setObjectName("status-label")
        self._timer_label = QLabel("")
        self._timer_label.setObjectName("timer-label")
        status_row.addWidget(self._status_label)
        status_row.addStretch()
        status_row.addWidget(self._timer_label)
        vl.addLayout(status_row)

        self._progress_bar = QProgressBar()
        self._progress_bar.setRange(0, 100)
        self._progress_bar.setValue(0)
        self._progress_bar.setFixedHeight(5)
        self._progress_bar.setTextVisible(False)
        vl.addWidget(self._progress_bar)

        return card

    # ── Start / Stop controls ─────────────────────────────────────────────────
    def _make_controls(self):
        w  = QWidget()
        w.setObjectName("root")
        hl = QHBoxLayout(w)
        hl.setContentsMargins(0, 0, 0, 0)
        hl.setSpacing(8)

        self._btn_start = QPushButton(self._s("btn_start"))
        self._btn_start.setObjectName("btn-primary")
        self._btn_start.setMinimumHeight(38)
        self._btn_start.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self._btn_start.clicked.connect(self._start_process)

        self._btn_stop = QPushButton(self._s("btn_stop"))
        self._btn_stop.setObjectName("btn-danger")
        self._btn_stop.setMinimumHeight(38)
        self._btn_stop.setEnabled(False)
        self._btn_stop.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self._btn_stop.clicked.connect(self._cancel_process)

        hl.addWidget(self._btn_start)
        hl.addWidget(self._btn_stop)
        return w

    # ── Log panel ─────────────────────────────────────────────────────────────
    def _make_log_panel(self):
        card   = QFrame()
        card.setObjectName("card")
        layout = QVBoxLayout(card)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        log_hdr = QFrame()
        log_hdr.setObjectName("log-header")
        log_hdr.setFixedHeight(36)
        hdr_layout = QHBoxLayout(log_hdr)
        hdr_layout.setContentsMargins(14, 0, 10, 0)

        self._lbl_log_header = QLabel(self._s("sec_log"))
        self._lbl_log_header.setObjectName("section-header")
        hdr_layout.addWidget(self._lbl_log_header)
        hdr_layout.addStretch()

        self._log_view = QPlainTextEdit()
        self._log_view.setObjectName("log")
        self._log_view.setReadOnly(True)

        self._btn_clear = QPushButton(self._s("btn_clear"))
        self._btn_clear.setFixedHeight(22)
        self._btn_clear.setMinimumWidth(52)
        self._btn_clear.setObjectName("btn-header")
        self._btn_clear.clicked.connect(self._log_view.clear if hasattr(self, "_log_view") else lambda: None)
        hdr_layout.addWidget(self._btn_clear)
        layout.addWidget(log_hdr)

        self._btn_clear.clicked.disconnect()
        self._btn_clear.clicked.connect(self._log_view.clear)
        layout.addWidget(self._log_view)

        return card

    # ── Footer ────────────────────────────────────────────────────────────────
    def _make_footer(self):
        footer = QFrame()
        footer.setObjectName("footer")
        footer.setFixedHeight(28)
        hl = QHBoxLayout(footer)
        hl.setContentsMargins(16, 0, 16, 0)
        hl.setSpacing(0)

        lbl_author = QLabel("acrilot")
        lbl_author.setObjectName("link")
        lbl_author.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        lbl_author.mousePressEvent = lambda e: QDesktopServices.openUrl(
            QUrl(GITHUB_PAGE_URL)
        )

        lbl_sep  = QLabel(" | ")
        lbl_sep.setObjectName("footer-text")
        self._lbl_footer_name = QLabel(f"Whisper GUI V{APP_VERSION}")
        self._lbl_footer_name.setObjectName("footer-text")

        hl.addWidget(lbl_author)
        hl.addWidget(lbl_sep)
        hl.addWidget(self._lbl_footer_name)
        hl.addStretch()
        return footer

    # ── Theme & language system ───────────────────────────────────────────────
    def _apply_theme(self, animate=True):
        """Rebuild and apply the full stylesheet + update accent-colored labels with smooth transition."""
        
        if animate:
            window = self.window()
            overlay = QLabel(window)
            pixmap = window.grab()
            overlay.setPixmap(pixmap)
            overlay.setGeometry(window.rect())
            overlay.show()

        ss = build_stylesheet(self._theme_name, self._dark_mode)
        QApplication.instance().setStyleSheet(ss)

        accent = THEME_ACCENTS[self._theme_name]["accent"]
        dim    = DARK_PALETTE["text_dim"] if self._dark_mode else LIGHT_PALETTE["text_dim"]
        txt    = DARK_PALETTE["text_main"] if self._dark_mode else LIGHT_PALETTE["text_main"]

        self._lbl_title.setStyleSheet(
            f"color:{accent}; font-size:19px; font-weight:700; letter-spacing:-0.5px;"
        )
        self._lbl_subtitle.setStyleSheet(f"color:{dim}; font-size:14px;")
        self._lbl_ver.setStyleSheet(f"color:{dim}; font-size:11px; font-family:monospace;")

        self.titlebar.title_label.setStyleSheet(f"color:{txt}; font-weight: 700; font-size: 13px;")

        border_color = '#ffffff' if self._dark_mode else '#000000'
        
        for name, btn in self._swatch_btns.items():
            color = THEME_ACCENTS[name]["accent"]
            if name == self._theme_name:
                btn.setStyleSheet(
                    f"QPushButton {{ background-color:{color}; border-radius:7px; border:2px solid {border_color}; }}"
                )
            else:
                btn.setStyleSheet(
                    f"QPushButton {{ background-color:{color}; border-radius:7px; border:none; }}"
                    f"QPushButton:hover {{ border:2px solid {txt}; }}"
                )

        if animate:
            QApplication.processEvents()
            effect = QGraphicsOpacityEffect(overlay)
            overlay.setGraphicsEffect(effect)
            self._theme_anim = QPropertyAnimation(effect, b"opacity")
            self._theme_anim.setDuration(500)
            self._theme_anim.setStartValue(1.0)
            self._theme_anim.setEndValue(0.0)
            self._theme_anim.setEasingCurve(QEasingCurve.Type.InOutQuad)
            self._theme_anim.finished.connect(overlay.deleteLater)
            self._theme_anim.start()

    def _apply_language(self):
        """Update all stored text labels to the current UI language."""
        S = STRINGS[self._ui_lang]

        self.setWindowTitle(S["window_title"])
        self.titlebar.title_label.setText(S["window_title"]) # Title bar text
        self._lbl_subtitle.setText(S["header_sub"])

        # Section headers
        for key, lbl in self._section_labels.items():
            lbl.setText(S[key].upper())

        # Field labels
        self._lbl_lang_field.setText(S["lbl_language"])
        self._lbl_compute.setText(S["lbl_compute"])
        self._lbl_beam.setText(S["lbl_beam"])

        # Inputs
        self._file_display.setPlaceholderText(S["ph_files"])
        self._output_display.setPlaceholderText(S["ph_output"])

        # Buttons
        self._btn_browse.setText(S["btn_browse"])
        self._btn_start.setText(S["btn_start"])
        self._btn_stop.setText(S["btn_stop"])
        self._btn_clear.setText(S["btn_clear"])

        # Checkboxes
        self._chk_vad.setText(S["chk_vad"])
        self._chk_single.setText(S["chk_single"])
        self._chk_timestamps.setText(S["chk_timestamps"])
        self._chk_translate.setText(S["chk_translate"])
        self._chk_install.setText(S["chk_install"])

        # Log header
        self._lbl_log_header.setText(S["sec_log"])

        # Language combo — preserve selected index
        idx = self._lang_combo.currentIndex()
        self._lang_combo.blockSignals(True)
        self._lang_combo.clear()
        self._lang_combo.addItems(S["lang_names"])
        self._lang_combo.setCurrentIndex(idx)
        self._lang_combo.blockSignals(False)

        # Status label (only update if showing "ready" equivalent)
        current_status = self._status_label.text()
        for other_lang in STRINGS.keys():
            if current_status == STRINGS[other_lang]["status_ready"]:
                self._status_label.setText(S["status_ready"])
                break

    def _change_language(self, text):
        # Integrates the language selected in the UI to the code
        self._ui_lang = self._ui_langs_map.get(text, "en")
        self._apply_language()
        self._save_preferences()

    def _toggle_mode(self):
        self._dark_mode = not self._dark_mode
        self._btn_mode.setText("☾" if self._dark_mode else "☀")
        self._apply_theme(animate=True)
        self._save_preferences()

    def _set_theme_color(self, name: str):
        self._theme_name = name
        self._apply_theme(animate=True)
        self._save_preferences()

    # ── Model scanning ────────────────────────────────────────────────────────
    def _scan_models(self):
        self._model_map = {}
        self._model_combo.blockSignals(True)
        self._model_combo.clear()
        default_idx = 0

        for i, (name, key) in enumerate(MODEL_OPTIONS.items()):
            icon    = "✓" if os.path.exists(MODEL_PATHS[key]) else "⬇"
            label   = f"{icon}  {name}"
            self._model_combo.addItem(label)
            self._model_map[label] = key
            if ("Faster Whisper" in name and "Large" in name
                    and os.path.exists(MODEL_PATHS[key])):
                default_idx = i

        self._model_combo.blockSignals(False)
        self._model_combo.setCurrentIndex(default_idx)
        self._on_model_change()

    def _on_model_change(self):
        self._update_output_path()
        key    = self._model_map.get(self._model_combo.currentText(), "")
        is_std = key.startswith("std_")
        self._compute_combo.setEnabled(not is_std)
        self._btn_stop.setVisible(not is_std)

    # ── File selection ────────────────────────────────────────────────────────
    def _select_files(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, "Select Audio or Video Files", "",
            "Media Files (*.mp3 *.wav *.m4a *.mp4 *.mkv *.flac *.ogg *.webm *.opus)"
            ";;All Files (*.*)",
        )
        if not files:
            return
        self._selected_files = files
        if len(files) == 1:
            self._file_display.setText(files[0])
            self._output_display.setEnabled(True)
            self._btn_out.setEnabled(True)
            self._update_output_path(files[0])
        else:
            self._file_display.setText(
                f"{len(files)} files selected (Batch Mode)"
            )
            folder = os.path.dirname(files[0])
            self._output_display.setText(folder + "  (batch)")
            self._output_display.setEnabled(False)
            self._btn_out.setEnabled(False)

    def _select_output(self):
        if len(self._selected_files) > 1:
            QMessageBox.information(self, "Info", self._s("batch_output_info"))
            return
        fmt     = self._format_combo.currentText()
        ext_map = {
            "PDF":  (".pdf",  "PDF Files (*.pdf)"),
            "DOCX": (".docx", "Word Files (*.docx)"),
            "XML":  (".xml",  "XML Subtitle (*.xml)"),
        }
        ext, flt = next(
            ((e, f) for k, (e, f) in ext_map.items() if k in fmt),
            (".txt", "Text Files (*.txt)"),
        )
        path, _ = QFileDialog.getSaveFileName(self, "Save Output", "", flt)
        if path:
            self._output_display.setText(path)

    def _on_format_change(self):
        if len(self._selected_files) == 1:
            self._update_output_path(self._selected_files[0])

    def _update_output_path(self, source_file=None):
        if source_file is None:
            if len(self._selected_files) == 1:
                source_file = self._selected_files[0]
            else:
                return
        folder, name  = os.path.split(source_file)
        stem          = os.path.splitext(name)[0]
        fmt           = self._format_combo.currentText()
        ext_map       = {"PDF": ".pdf", "DOCX": ".docx", "XML": ".xml"}
        ext           = next((v for k, v in ext_map.items() if k in fmt), ".txt")
        key           = self._model_map.get(self._model_combo.currentText(), "whisper")
        trans_suffix  = "_translated" if self._chk_translate.isChecked() else ""
        fname         = f"{stem}_{key.replace('_','-')}{trans_suffix}_transcript{ext}"
        self._output_display.setText(os.path.join(folder, fname))

    # ── Process control ───────────────────────────────────────────────────────
    def _start_process(self):
        if not self._selected_files:
            QMessageBox.warning(self, "Warning", self._s("warn_no_file"))
            return

        params = {
            "install_mode":   self._chk_install.isChecked(),
            "model_display":  self._model_combo.currentText(),
            "model_map":      dict(self._model_map),
            "language_index": self._lang_combo.currentIndex(),
            "compute_type":   self._compute_combo.currentText(),
            "beam_size":      self._beam_spin.value(),
            "vad":            self._chk_vad.isChecked(),
            "single_block":   self._chk_single.isChecked(),
            "timestamps":     self._chk_timestamps.isChecked(),
            "translate":      self._chk_translate.isChecked(),
            "output_format":  self._format_combo.currentText(),
            "files":          list(self._selected_files),
            "output_path":    self._output_display.text(),
        }

        self._log_view.clear()
        self._progress_bar.setValue(0)
        self._btn_start.setEnabled(False)
        self._btn_stop.setEnabled(True)
        self._start_timer()

        self._worker = TranscriptWorker(params, self._model_ref, self._model_lock)
        self._worker.log.connect(self._append_log)
        self._worker.progress.connect(lambda v: self._progress_bar.setValue(int(v)))
        self._worker.status.connect(self._status_label.setText)
        self._worker.done.connect(self._on_done)
        self._worker.error.connect(self._on_error)
        self._worker.start()

    def _cancel_process(self):
        reply = QMessageBox.question(
            self, self._s("stop_title"), self._s("stop_msg"),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if reply == QMessageBox.StandardButton.Yes and self._worker:
            self._worker.request_cancel()
            self._btn_stop.setEnabled(False)

    def _on_done(self, success, elapsed):
        self._stop_timer()
        self._btn_start.setEnabled(True)
        self._btn_stop.setEnabled(False)
        if success:
            self._progress_bar.setValue(100)
            self._status_label.setText(self._s("status_done"))
            QMessageBox.information(
                self, self._s("done_title"),
                self._s("done_msg").format(elapsed),
            )
        else:
            self._status_label.setText(self._s("status_stopped"))

    def _on_error(self, title, message):
        self._stop_timer()
        self._btn_start.setEnabled(True)
        self._btn_stop.setEnabled(False)
        self._status_label.setText(self._s("status_error"))
        QMessageBox.critical(self, title, message)

    def _append_log(self, text):
        self._log_view.appendPlainText(text)
        sb = self._log_view.verticalScrollBar()
        sb.setValue(sb.maximum())

    # ── Timer ─────────────────────────────────────────────────────────────────
    def _setup_timer(self):
        self._qtimer = QTimer(self)
        self._qtimer.setInterval(1000)
        self._qtimer.timeout.connect(self._tick_timer)

    def _start_timer(self):
        self._timer_start  = time.time()
        self._timer_active = True
        self._qtimer.start()

    def _stop_timer(self):
        self._timer_active = False
        self._qtimer.stop()
        self._timer_label.setText("")

    def _tick_timer(self):
        if not self._timer_active:
            return
        elapsed = int(time.time() - self._timer_start)
        m, s    = divmod(elapsed, 60)
        self._timer_label.setText(f"{m}:{s:02d}")
        
    # ── Update system ─────────────────────────────────────────────────────────
    def _on_update_available(self, new_version, asset):
        msg = self._s("update_msg").format(new=new_version, cur=APP_VERSION)
        reply = QMessageBox.question(
            self, self._s("update_title"), msg,
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if reply != QMessageBox.StandardButton.Yes:
            return

        if asset is None:
            QMessageBox.warning(self, self._s("update_title"),
                                self._s("update_no_asset"))
            QDesktopServices.openUrl(
                QUrl(f"https://github.com/{GITHUB_REPO}/releases/latest")
            )
            return

        target = os.path.abspath(sys.argv[0])
        backup = target + ".bak"
        tmp    = target + ".tmp"

        try:
            urllib.request.urlretrieve(asset["browser_download_url"], tmp)
            if os.path.exists(backup):
                os.remove(backup)
            os.rename(target, backup)
            os.rename(tmp, target)
            QMessageBox.information(
                self, self._s("update_title"),
                self._s("update_done").format(new=new_version),
            )
            subprocess.Popen([sys.executable] + sys.argv)
            sys.exit(0)
        except Exception as e:
            if os.path.exists(backup) and not os.path.exists(target):
                try:
                    os.rename(backup, target)
                except OSError:
                    pass
            if os.path.exists(tmp):
                try:
                    os.remove(tmp)
                except OSError:
                    pass
            QMessageBox.critical(
                self, self._s("update_title"),
                self._s("update_fail").format(err=e),
            )

    # ── Close event ───────────────────────────────────────────────────────────
    def closeEvent(self, event):
        if self._worker and self._worker.isRunning():
            reply = QMessageBox.question(
                self, self._s("exit_title"), self._s("exit_msg"),
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            )
            if reply != QMessageBox.StandardButton.Yes:
                event.ignore()
                return
            self._worker.request_cancel()
            self._worker.wait(3000)

        with self._model_lock:
            m = self._model_ref.get("model")
            if m:
                del m
                self._model_ref["model"] = None
                gc.collect()
                try:
                    import torch
                    if torch.cuda.is_available():
                        torch.cuda.empty_cache()
                except ImportError:
                    pass

        event.accept()


# ── Preferences System ────────────────────────────────────────────────────────
    def _get_config_path(self):
        return os.path.join(BASE_DIR, "whisper_gui_config.json")

    def _load_preferences(self):
        config_path = self._get_config_path()
        if os.path.exists(config_path):
            try:
                with open(config_path, "r", encoding="utf-8") as f:
                    config = json.load(f)
                    self._ui_lang = config.get("ui_lang", self._ui_lang)
                    self._theme_name = config.get("theme_name", self._theme_name)
                    self._dark_mode = config.get("dark_mode", self._dark_mode)
            except Exception:
                pass

    def _save_preferences(self):
        config_path = self._get_config_path()
        try:
            with open(config_path, "w", encoding="utf-8") as f:
                json.dump({
                    "ui_lang": self._ui_lang,
                    "theme_name": self._theme_name,
                    "dark_mode": self._dark_mode
                }, f, indent=4)
        except Exception:
            pass


# ── Bootstrap installer ───────────────────────────────────────────────────────
def _bootstrap_install():
    import tkinter as tk
    from tkinter import messagebox

    root = tk.Tk()
    root.withdraw()

    if not messagebox.askyesno(
        "PyQt6 Required",
        "PyQt6 is not installed. It is required to run this application.\n\n"
        "Install all required libraries now?\n"
        "(This includes PyQt6, PyTorch, Whisper, and other dependencies.)",
    ):
        root.destroy()
        return

    root.destroy()

    install_win = tk.Tk()
    install_win.title("Installing...")
    install_win.geometry("420x80")
    install_win.resizable(False, False)
    tk.Label(
        install_win,
        text="Installing PyQt6 and dependencies, please wait...",
        font=("Segoe UI", 10),
        pady=24,
    ).pack()
    install_win.update()

    try:
        for lib in LIBS_TO_INSTALL:
            subprocess.run(
                [sys.executable, "-m", "pip", "install", lib],
                check=True, capture_output=True,
            )
        install_win.destroy()
        done_root = tk.Tk()
        done_root.withdraw()
        messagebox.showinfo("Done",
            "Installation complete. The application will now restart.")
        done_root.destroy()
        subprocess.Popen([sys.executable] + sys.argv)
    except Exception as e:
        install_win.destroy()
        err_root = tk.Tk()
        err_root.withdraw()
        messagebox.showerror("Error", f"Installation failed:\n{e}")
        err_root.destroy()

    sys.exit(0)

# ── Entry point ───────────────────────────────────────────────────────────────
def main():
    if not PYQT6_OK:
        _bootstrap_install()
        return

    app = QApplication(sys.argv)
    app.setApplicationName("Whisper GUI")
    app.setStyle("WindowsVista")

    window = WhisperApp()
    window.show()
    window.showMaximized()

    sys.exit(app.exec())

if __name__ == "__main__":
    main()
